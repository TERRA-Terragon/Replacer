# подключение библиотек

# работа с файлами ".docx"
import docx
# в данном приложении используется для кроссплатформенной работы с путями файлов
import os
# в данном приложении используется для копии из выбранного пользователем файла
import shutil
# работа с GUI (графический интерфейс пользователя)
import tkinter as tk
from tkinter import filedialog, scrolledtext
# подключение класса "Replacer" с функциональностью для замены текста из файла "replacer.py"
from replacer import Replacer

replacer = Replacer()




# функция для чтения всего файла, принимает на вход название файла
def read_txt_file(path: str) -> str:
    # открытие файла
    with open(path, "r", encoding="utf-8") as f:
        # чтение файла
        return f.read()

# функция для чтения всего docx файла, принимает на вход название docx файла
def read_docx_file(path: str) -> str:
    # открытие файла
    document = docx.Document(path)
    # чтение файла
    return "\n".join(p.text for p in document.paragraphs)

# функция для обработки входного файла
def process_file(input_path: str, text_widget: tk.Text):
    try:
        # проверка формата входного файла
        ext = os.path.splitext(input_path)[1].lower()
        if ext == ".txt":
            # если .txt чтение как .txt файл
            text = read_txt_file(input_path)
        elif ext == ".docx":
            # если .docx чтение как .docx файл
            text = read_docx_file(input_path)
        else:
            # если формат файла иной -> ошибка в пользовательском окне
            text_widget.insert(tk.END, f"Формат файла не поддерживается: {ext}\n\
            Поддерживаемые: '.docx', '.txt'")
            return

        # создание нового файла с "_copy" в названии (file.txt -> file_copy.txt)
        output_path = os.path.splitext(input_path)[0] + "_copy" + ext

        # копирование оригинала в новый файл
        shutil.copy(input_path, output_path)

        # обработка содержимого
        processed_text = replacer.replace_all(text)

        # сохранение в новый файл
        if ext == ".txt":
            with open(output_path, "w", encoding="utf-8") as f:
                f.write(processed_text)
        elif ext == ".docx":
            doc = docx.Document()
            for line in processed_text.splitlines():
                doc.add_paragraph(line)
            doc.save(output_path)

        # вывод в окне пользователя об успехе
        text_widget.insert(tk.END, f"Успешно обработано: {output_path}\n")

    # обработка ошибки
    except Exception as e:
        text_widget.insert(tk.END, f"Ошибка при обработке: {e}\n")


# обработчик выбора файла, будет срабатывать каждый раз при нажатии на кнопку выбора файла в интерфейсе
def select_file(text_widget: tk.Text):
    # открытие файлового менеджера и сохранение пути к выбранному пользователем файлу
    file_path = filedialog.askopenfilename(title="Выберите файл", filetypes=[("TXT and DOCX", "*.txt *.docx")])
    # если пользователь что-то выбрал
    if file_path:
        # вывод уведомления о выбранном файле и запуск его обработки
        text_widget.insert(tk.END, f"Выбран файл: {file_path}\n")
        process_file(file_path, text_widget)

# основная функция программы
def main():
    # функция для создания пользовательского окна
    def run_gui():
        # создение окна размером 600x400 пикселей
        root = tk.Tk()
        root.title("Data Replacer")
        root.geometry("600x400")

        # кнопка выбора файла
        btn = tk.Button(root, text="Выбрать файл", command=lambda: select_file(text_area))
        btn.pack(pady=10)

        # текстовое поле для вывода логов
        text_area = scrolledtext.ScrolledText(root, wrap=tk.WORD)
        text_area.pack(expand=True, fill=tk.BOTH, padx=10, pady=10)

        # зацкиливание программы чтоб она не завершилась сама по себе
        root.mainloop()

    run_gui()

# запуск основной функции программы
if __name__ == "__main__":
    main()